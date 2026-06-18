from docx import Document


class DOCXExtractor:
    def extract(self, file_path: str) -> str:
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    tables_text.append(row_text)
        return "\n\n".join(paragraphs + tables_text)
